import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from langdetect import detect
from googletrans import Translator
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import fitz
import os


def detectar_idioma(texto):
    try:
        return detect(texto)
    except:
        return "Desconhecido"

def traduzir_texto(texto, destino, translator):
    try:
        return translator.translate(texto, dest=destino).text
    except Exception as e:
        print("Erro ao traduzir:", e)
        return texto

def extrair_texto_pdf(caminho_pdf):
    texto_total = ""
    with fitz.open(caminho_pdf) as doc:
        for pagina in doc:
            texto_total += pagina.get_text()
    return texto_total

def salvar_texto_em_pdf(texto, caminho_saida):
    c = canvas.Canvas(caminho_saida, pagesize=A4)
    largura, altura = A4
    y = altura - 50
    linhas = texto.split("\n")
    for linha in linhas:
        if y < 50:  # nova página
            c.showPage()
            y = altura - 50
        c.drawString(40, y, linha)
        y -= 15
    c.save()

def traduzir_pdf(caminho_pdf, idioma_destino):
    texto = extrair_texto_pdf(caminho_pdf)
    idioma_original = detectar_idioma(texto)
    translator = Translator()
    texto_traduzido = traduzir_texto(texto, idioma_destino, translator)

    nome_saida = os.path.splitext(caminho_pdf)[0] + f"_traduzido_para_{idioma_destino}.pdf"
    salvar_texto_em_pdf(texto_traduzido, nome_saida)

    return nome_saida, idioma_original

def traduzir_docx(caminho_docx, idioma_destino):
    doc = Document(caminho_docx)
    translator = Translator()
    texto_completo = "\n".join([p.text for p in doc.paragraphs])
    idioma_original = detectar_idioma(texto_completo)

    novo_doc = Document()

    for paragrafo in doc.paragraphs:
        novo_paragrafo = novo_doc.add_paragraph()

        for run in paragrafo.runs:
            texto_original = run.text.strip()
            if texto_original:
                texto_traduzido = traduzir_texto(texto_original, idioma_destino, translator)
                novo_run = novo_paragrafo.add_run(texto_traduzido)

                # Preservar estilo
                novo_run.bold = run.bold
                novo_run.italic = run.italic
                novo_run.underline = run.underline
                novo_run.font.size = run.font.size
                novo_run.font.name = run.font.name
            else:
                novo_paragrafo.add_run("")

    nome_saida = os.path.splitext(caminho_docx)[0] + f"_traduzido_para_{idioma_destino}.docx"
    novo_doc.save(nome_saida)
    return nome_saida, idioma_original

# Interface gráfica
def selecionar_arquivo():
    caminho = filedialog.askopenfilename(filetypes=[("Documentos Word ou PDF", "*.docx *.pdf")])
    if caminho:
        entrada_path.set(caminho)

def iniciar_traducao():
    arquivo = entrada_path.get()
    idioma = idioma_destino.get()

    if not arquivo or not idioma:
        messagebox.showwarning("Aviso", "Selecione um arquivo e um idioma.")
        return

    try:
        if arquivo.endswith(".docx"):
            saida, idioma_detectado = traduzir_docx(arquivo, idioma)
        elif arquivo.endswith(".pdf"):
            saida, idioma_detectado = traduzir_pdf(arquivo, idioma)
        else:
            raise ValueError("Formato de arquivo não suportado.")
        messagebox.showinfo("Sucesso", f"Idioma detectado: {idioma_detectado}\nArquivo salvo: {saida}")
    except Exception as e:
        messagebox.showerror("Erro", str(e))

# GUI
janela = tk.Tk()
janela.title("Tradutor de Documentos (DOCX / PDF)")
janela.geometry("400x250")

entrada_path = tk.StringVar()
idioma_destino = tk.StringVar()

tk.Label(janela, text="Selecione o arquivo .docx ou .pdf:").pack(pady=5)
tk.Entry(janela, textvariable=entrada_path, width=50).pack()
tk.Button(janela, text="Procurar", command=selecionar_arquivo).pack(pady=5)

tk.Label(janela, text="Selecione o idioma de destino:").pack(pady=10)
idiomas = {
    "Português": "pt",
    "Inglês": "en",
    "Espanhol": "es",
    "Francês": "fr",
    "Alemão": "de",
    "Italiano": "it"
}
idioma_combo = ttk.Combobox(janela, values=list(idiomas.keys()), state="readonly")
idioma_combo.pack()
idioma_combo.bind("<<ComboboxSelected>>", lambda e: idioma_destino.set(idiomas[idioma_combo.get()]))

tk.Button(janela, text="Traduzir Documento", command=iniciar_traducao, bg="green", fg="white").pack(pady=20)

janela.mainloop()